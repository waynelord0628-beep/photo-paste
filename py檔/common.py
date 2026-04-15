"""
共用工具模組 - 供所有圖片彙整腳本使用
"""

import sys
import os
import re
import io
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

# ── HEIC 支援（選用）────────────────────────────────────────────
# 若已安裝 pillow-heif，自動啟用 HEIC/HEIF 支援
try:
    import pillow_heif

    pillow_heif.register_heif_opener()
    _HEIC_SUPPORTED = True
except ImportError:
    _HEIC_SUPPORTED = False

from PIL import Image


# ── 路徑工具 ────────────────────────────────────────────────────


def get_base_path():
    """
    取得資源（模板檔等）所在資料夾路徑，相容 PyInstaller .exe。
    - 打包後（frozen）：PyInstaller 將 --add-data 資源解壓到 sys._MEIPASS，
      模板就在該目錄根層，故回傳 sys._MEIPASS。
    - 開發時直接跑 .py：模板放在 py檔 的上一層資料夾，
      故回傳 __file__ 所在目錄的上一層。
    """
    if getattr(sys, "frozen", False):
        return sys._MEIPASS  # type: ignore[attr-defined]
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def get_unique_filename(folder, base_name, ext):
    """若檔名已存在，自動在後方加序號避免覆蓋"""
    filename = f"{base_name}{ext}"
    counter = 1
    while os.path.exists(os.path.join(folder, filename)):
        filename = f"{base_name}_{counter}{ext}"
        counter += 1
    return filename


# ── 圖片清單 ────────────────────────────────────────────────────

# python-docx 可直接插入的格式
_DOCX_NATIVE_EXTS = {".jpg", ".jpeg", ".png", ".gif", ".bmp"}

# 需要透過 Pillow 轉換才能插入的格式
_PILLOW_CONVERT_EXTS = {".webp", ".tiff", ".tif"}

# HEIC/HEIF（需要 pillow-heif）
_HEIC_EXTS = {".heic", ".heif"}


def _build_valid_exts():
    """動態建立支援的副檔名集合（不分大小寫）"""
    exts = _DOCX_NATIVE_EXTS | _PILLOW_CONVERT_EXTS
    if _HEIC_SUPPORTED:
        exts = exts | _HEIC_EXTS
    # 同時加入大寫版本
    return exts | {e.upper() for e in exts}


VALID_IMAGE_EXTS = _build_valid_exts()


def load_images(path_captures):
    """
    從指定資料夾讀取合法圖片，回傳 (完整路徑列表, 無副檔名名稱列表)。
    若資料夾不存在會 raise FileNotFoundError。
    """
    if not os.path.exists(path_captures):
        raise FileNotFoundError(f"找不到 照片放這 資料夾，請確認路徑：{path_captures}")
    names = sorted(
        fn
        for fn in os.listdir(path_captures)
        if os.path.splitext(fn)[1].lower() in {e.lower() for e in VALID_IMAGE_EXTS}
    )
    paths = [os.path.join(path_captures, fn) for fn in names]
    names_noext = [os.path.splitext(fn)[0] for fn in names]
    return paths, names_noext


# ── 圖片插入輔助 ─────────────────────────────────────────────────


def open_image_as_stream(img_path):
    """
    將圖片轉成 python-docx 可插入的 bytes stream（PNG 格式）。
    - JPEG/PNG：直接讀取原始檔案（保持品質）
    - HEIC、WEBP、TIFF 等其他格式：用 Pillow 轉成 PNG bytes
    """
    ext = os.path.splitext(img_path)[1].lower()
    if ext in (".jpg", ".jpeg", ".png"):
        # 原生支援，直接回傳檔案路徑即可（讓 add_picture 直接讀檔）
        return img_path
    # 其他格式：用 Pillow 讀取後輸出成 PNG bytes
    with Image.open(img_path) as img:
        # RGBA 轉 RGB（Word 不支援帶透明通道的 PNG 在某些情況下）
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGBA")
        else:
            img = img.convert("RGB")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        return buf


# ── Word 文件操作 ────────────────────────────────────────────────


def open_template(template_path):
    """開啟 Word 模板，若找不到則 raise FileNotFoundError"""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"找不到模板檔，請確認路徑：{template_path}")
    return docx.Document(template_path)


def setup_header(document, section, title_text):
    """
    設定頁首：距離、標題文字（標楷體 20pt 置中）、下方空行。
    """
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(0.3)

    header = section.header
    if header.paragraphs:
        p_header = header.paragraphs[0]
        p_header.clear()
    else:
        p_header = header.add_paragraph()

    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.style = document.styles["Normal"]

    fmt = p_header.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(20)

    run = p_header.add_run(title_text)
    run.font.size = Pt(20)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "標楷體")

    header.add_paragraph("")


def set_run_font(
    run, size_pt=14, western_font="Times New Roman", east_asia_font="標楷體"
):
    """設定 run 的字體大小與中西文字型"""
    run.font.size = Pt(size_pt)
    run.font.name = western_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)


def delete_first_paragraph_if_empty(doc):
    """刪除文件開頭的空白段落（Word 模板常見殘留）"""
    if not doc.paragraphs:
        return
    first_p = doc.paragraphs[0]
    if not first_p.text.strip():
        elem = first_p._element
        elem.getparent().remove(elem)


def delete_trailing_empty_paragraphs(doc):
    """
    刪除文件結尾所有多餘的空白段落，防止 Word 產生額外空白頁。

    策略：從 body 最後一個子元素往前掃：
    - 若是空的 <w:p>（無文字、無圖片），且其 pPr 內沒有帶 sectPr
      （即不是分節段落），直接刪除。
    - 若是非空段落或表格，停止。
    - 若 <w:p> 內帶有 <w:pPr><w:sectPr>（分節符），
      把 sectPr 移到 body 直屬的最終 <w:sectPr>，再刪掉該段落，
      這樣 Word 就不會把它視為「下一頁起點」。
    """
    from docx.oxml.ns import qn as _qn
    from lxml import etree

    body = doc.element.body
    W_P = _qn("w:p")
    W_SECT_PR = _qn("w:sectPr")
    W_PPR = _qn("w:pPr")

    while True:
        children = list(body)
        if not children:
            break
        last = children[-1]

        # 最終的 <w:sectPr> 直接掛在 body 下，不動它
        if last.tag == W_SECT_PR:
            break

        # 只處理 <w:p>
        if last.tag != W_P:
            break

        # 若段落有實際文字或圖片內容，停止
        text = "".join(t.text or "" for t in last.iter(_qn("w:t"))).strip()
        has_drawing = last.find(".//" + _qn("w:drawing")) is not None
        if text or has_drawing:
            break

        # 若段落內的 pPr 帶有 sectPr（分節符），搬移到 body 的最終 sectPr
        pPr = last.find(W_PPR)
        if pPr is not None:
            inline_sect = pPr.find(W_SECT_PR)
            if inline_sect is not None:
                # 取得或建立 body 最終 sectPr
                body_sect = body.find(W_SECT_PR)
                if body_sect is None:
                    body_sect = etree.SubElement(body, W_SECT_PR)
                # 把分節資訊的子元素搬過去（避免覆蓋已有設定）
                for child in list(inline_sect):
                    body_sect.append(child)
                pPr.remove(inline_sect)

        body.remove(last)


# ── 儲存格寬度工具 ────────────────────────────────────────────────


def set_cell_width(cell, width_cm):
    """
    強制設定 python-docx 儲存格的寬度（公分）。
    python-docx 的 cell.width 不可靠，需直接操作 XML tcPr/tcW。
    """
    from lxml import etree
    from docx.oxml.ns import qn as _qn
    from docx.shared import Cm

    width_emu = int(Cm(width_cm))  # Cm 物件即 EMU
    # Word 表格寬度單位是 twips（1 inch = 1440 twips，1 cm ≈ 567 twips）
    twips = int(width_emu / 914400 * 1440)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 移除舊的 tcW（若存在）
    for old in tcPr.findall(_qn("w:tcW")):
        tcPr.remove(old)
    tcW = etree.SubElement(tcPr, _qn("w:tcW"))
    tcW.set(_qn("w:w"), str(twips))
    tcW.set(_qn("w:type"), "dxa")


def fill_name_cell(
    cell,
    number: int,
    desc_text: str,
    outer_width_cm: float = 9.0,
    num_col_cm: float = 2.2,
):
    """
    將名稱格（單一 cell）拆成左右兩個巢狀欄位：
      左格（固定 num_col_cm cm）：「編號 N」，置中
      右格（其餘寬度）          ：desc_text，靠左

    outer_width_cm：外層 cell 的實際欄寬（cm），用來計算右格寬度。
    做法：在 cell 內插入一個 1 列 2 欄的巢狀表格，
    外層 cell 本身清空，讓巢狀表格填滿。
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn
    from lxml import etree

    # 清除 cell 原有段落內容
    tc = cell._tc
    for p in tc.findall(_qn("w:p")):
        tc.remove(p)

    # 取得外層 cell 的表格物件，以便用 add_table
    # python-docx 不支援直接在 cell 內 add_table，改用 XML 直接建
    # ── 建立巢狀表格 XML ──
    tbl_xml = etree.SubElement(tc, _qn("w:tbl"))

    # tblPr
    tblPr = etree.SubElement(tbl_xml, _qn("w:tblPr"))
    tblStyle = etree.SubElement(tblPr, _qn("w:tblStyle"))
    tblStyle.set(_qn("w:val"), "TableGrid")
    tblW = etree.SubElement(tblPr, _qn("w:tblW"))
    tblW.set(_qn("w:w"), "0")
    tblW.set(_qn("w:type"), "auto")
    tblLayout = etree.SubElement(tblPr, _qn("w:tblLayout"))
    tblLayout.set(_qn("w:type"), "fixed")
    # 移除巢狀表格外框，但保留 insideV（左右格之間的中間直格線）
    tblBorders = etree.SubElement(tblPr, _qn("w:tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH"):
        b = etree.SubElement(tblBorders, _qn(f"w:{side}"))
        b.set(_qn("w:val"), "none")
    insideV = etree.SubElement(tblBorders, _qn("w:insideV"))
    insideV.set(_qn("w:val"), "single")
    insideV.set(_qn("w:sz"), "4")
    insideV.set(_qn("w:space"), "0")
    insideV.set(_qn("w:color"), "000000")

    # tblGrid（兩欄）
    tblGrid = etree.SubElement(tbl_xml, _qn("w:tblGrid"))
    # 左格寬度（twips）
    num_twips = int(num_col_cm / 2.54 * 1440)
    # 右格寬度 = 外層欄寬 - 左格寬度
    outer_twips = int(outer_width_cm / 2.54 * 1440)
    desc_twips = max(outer_twips - num_twips, 500)
    gridCol1 = etree.SubElement(tblGrid, _qn("w:gridCol"))
    gridCol1.set(_qn("w:w"), str(num_twips))
    gridCol2 = etree.SubElement(tblGrid, _qn("w:gridCol"))
    gridCol2.set(_qn("w:w"), str(desc_twips))

    # ── 單一列 ──
    tr = etree.SubElement(tbl_xml, _qn("w:tr"))
    trPr = etree.SubElement(tr, _qn("w:trPr"))
    trHeight = etree.SubElement(trPr, _qn("w:trHeight"))
    trHeight.set(_qn("w:val"), "0")
    trHeight.set(_qn("w:hRule"), "auto")

    def _make_tc(width_twips, text, align="center", fit_text=False):
        tc_el = etree.SubElement(tr, _qn("w:tc"))
        tcPr_el = etree.SubElement(tc_el, _qn("w:tcPr"))
        tcW_el = etree.SubElement(tcPr_el, _qn("w:tcW"))
        tcW_el.set(_qn("w:w"), str(width_twips))
        tcW_el.set(_qn("w:type"), "dxa")
        # 垂直置中
        vAlign = etree.SubElement(tcPr_el, _qn("w:vAlign"))
        vAlign.set(_qn("w:val"), "center")
        # 段落
        p_el = etree.SubElement(tc_el, _qn("w:p"))
        pPr_el = etree.SubElement(p_el, _qn("w:pPr"))
        jc = etree.SubElement(pPr_el, _qn("w:jc"))
        jc.set(_qn("w:val"), align)
        # 段落間距
        spacing = etree.SubElement(pPr_el, _qn("w:spacing"))
        spacing.set(_qn("w:before"), "0")
        spacing.set(_qn("w:after"), "0")
        r_el = etree.SubElement(p_el, _qn("w:r"))
        rPr_el = etree.SubElement(r_el, _qn("w:rPr"))
        # 字型 14pt 標楷體
        rFonts = etree.SubElement(rPr_el, _qn("w:rFonts"))
        rFonts.set(_qn("w:ascii"), "Times New Roman")
        rFonts.set(_qn("w:eastAsia"), "標楷體")
        sz = etree.SubElement(rPr_el, _qn("w:sz"))
        sz.set(_qn("w:val"), "28")  # 14pt = 28 half-points
        szCs = etree.SubElement(rPr_el, _qn("w:szCs"))
        szCs.set(_qn("w:val"), "28")
        # 文字自適應：超出欄寬時自動縮小字型
        if fit_text:
            fitText = etree.SubElement(rPr_el, _qn("w:fitText"))
            fitText.set(_qn("w:val"), str(width_twips))
            fitText.set(_qn("w:id"), "1")
        t_el = etree.SubElement(r_el, _qn("w:t"))
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = text
        return tc_el

    _make_tc(num_twips, f"編號 {number}", align="center", fit_text=True)
    _make_tc(desc_twips, desc_text, align="left", fit_text=True)

    # cell 最後需要一個段落（Word 規範）
    closing_p = etree.SubElement(tc, _qn("w:p"))


def set_table_fixed_width(tbl, total_width_cm):
    """
    強制設定表格的總寬度為固定（fixed）模式。
    避免 Word 在多欄表格中自動縮放欄寬。
    """
    from lxml import etree
    from docx.oxml.ns import qn as _qn
    from docx.shared import Cm

    total_emu = int(Cm(total_width_cm))
    twips = int(total_emu / 914400 * 1440)

    tblPr = tbl._tbl.find(_qn("w:tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl._tbl, _qn("w:tblPr"))

    # 設定 tblW = fixed
    for old in tblPr.findall(_qn("w:tblW")):
        tblPr.remove(old)
    tblW = etree.SubElement(tblPr, _qn("w:tblW"))
    tblW.set(_qn("w:w"), str(twips))
    tblW.set(_qn("w:type"), "dxa")

    # 設定 tblLayout = fixed
    for old in tblPr.findall(_qn("w:tblLayout")):
        tblPr.remove(old)
    tblLayout = etree.SubElement(tblPr, _qn("w:tblLayout"))
    tblLayout.set(_qn("w:type"), "fixed")


# ── 儲存文件 ────────────────────────────────────────────────────


def save_document(document, path_now, title_text):
    """
    將文件儲存到 path_now 目錄，檔名依標題產生（去除非法字元）。
    回傳最終輸出的完整路徑。
    """
    safe_title = re.sub(r'[\\/*?:"<>|]', "_", title_text)
    filename = get_unique_filename(path_now, safe_title, ".docx")
    output_path = os.path.join(path_now, filename)
    document.save(output_path)
    print(f"完成！輸出檔案：{output_path}")
    return output_path
