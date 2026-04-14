import os
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from common import (
    get_base_path,
    load_images,
    open_template,
    setup_header,
    set_run_font,
    open_image_as_stream,
    delete_first_paragraph_if_empty,
    save_document,
)

# ── 路徑設定 ────────────────────────────────────────────────────
path_now = get_base_path()
template_path = os.path.join(path_now, "word橫向模板別動.docx")
path_captures = os.path.join(path_now, "照片放這")

# ── 讀取圖片 ────────────────────────────────────────────────────
image_file_path, image_file_name_noext = load_images(path_captures)

# ── 開啟模板、輸入標題 ──────────────────────────────────────────
document = open_template(template_path)
title_text = input("請輸入要生成的標題：").strip()

# ── 設定頁首 ────────────────────────────────────────────────────
section = document.sections[0]
setup_header(document, section, title_text)

# ── 每頁 3 張橫式，含檔名 ────────────────────────────────────────
GROUP_SIZE = 3

for group_start in range(0, len(image_file_path), GROUP_SIZE):
    group_imgs = image_file_path[group_start : group_start + GROUP_SIZE]
    group_names = image_file_name_noext[group_start : group_start + GROUP_SIZE]

    tbl = document.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    tbl.rows[0].height = Cm(15)
    tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    tbl.rows[1].height = Cm(1.6)
    tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for i, (img_path, img_name) in enumerate(zip(group_imgs, group_names)):
        # 圖片列
        cell_pic = tbl.cell(0, i)
        cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_pic = cell_pic.paragraphs[0]
        p_pic.add_run().add_picture(open_image_as_stream(img_path), height=Cm(14))
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 檔名列
        cell_name = tbl.cell(1, i)
        cell_name.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_name = cell_name.paragraphs[0]
        run_name = p_name.add_run(img_name)
        set_run_font(run_name)
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if group_start + GROUP_SIZE < len(image_file_path):
        document.add_section(WD_SECTION.NEW_PAGE)

# ── 收尾、儲存 ──────────────────────────────────────────────────
delete_first_paragraph_if_empty(document)
save_document(document, path_now, title_text)
