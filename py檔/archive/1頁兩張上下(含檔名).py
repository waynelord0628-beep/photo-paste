import os
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from common import (
    get_base_path,
    load_images,
    open_template,
    setup_header,
    set_run_font,
    open_image_as_stream,
    delete_first_paragraph_if_empty,
    save_document,
)

# ── 路徑設定 ────────────────────────────────────────────────────
path_now = get_base_path()
template_path = os.path.join(path_now, "word模板別動.docx")
path_captures = os.path.join(path_now, "照片放這")

# ── 讀取圖片 ────────────────────────────────────────────────────
image_file_path, image_file_name_noext = load_images(path_captures)

# ── 開啟模板、設定頁邊界、輸入標題 ────────────────────────────────
document = open_template(template_path)

section = document.sections[0]
section.left_margin = Cm(1.27)
section.right_margin = Cm(1.27)
section.top_margin = Cm(1.27)
section.bottom_margin = Cm(1.27)

title_text = input("請輸入要生成的標題：").strip()

# ── 設定頁首 ────────────────────────────────────────────────────
setup_header(document, section, title_text)

# ── 建立表格：(圖片數 × 2) 列，1 欄 ─────────────────────────────
PIC_HEIGHT = Cm(14.5)  # 原始 3880000 EMU ≈ 13.76 cm；維持原始值
PIC_HEIGHT_EMU = 3880000

tbl = document.add_table(rows=len(image_file_path) * 2, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"

for i, (img_path, img_name) in enumerate(zip(image_file_path, image_file_name_noext)):
    # 圖片列
    cell_pic = tbl.cell(i * 2, 0)
    p_pic = cell_pic.paragraphs[0]
    p_pic.add_run().add_picture(open_image_as_stream(img_path), height=PIC_HEIGHT_EMU)
    p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 檔名列
    cell_name = tbl.cell(i * 2 + 1, 0)
    cell_name.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_name = cell_name.paragraphs[0]
    run_name = p_name.add_run(img_name)
    set_run_font(run_name)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = tbl.rows[i * 2 + 1]
    row.height = Pt(28.35)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

# ── 收尾、儲存 ──────────────────────────────────────────────────
delete_first_paragraph_if_empty(document)
save_document(document, path_now, title_text)
