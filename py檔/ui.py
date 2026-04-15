"""
圖片彙整工具 - PyQt6 圖形介面
"""

import os
import sys
import traceback
from PyQt6.QtCore import (
    Qt,
    QThread,
    pyqtSignal,
    QRect,
    QSize,
    QPropertyAnimation,
    QEasingCurve,
    QPoint as _QPoint,
    QTimer,
)
from PyQt6.QtWidgets import (
    QApplication,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QGridLayout,
    QLabel,
    QLineEdit,
    QPushButton,
    QButtonGroup,
    QTextEdit,
    QMessageBox,
    QFileDialog,
    QFrame,
    QSizePolicy,
    QScrollArea,
    QGraphicsOpacityEffect,
    QGraphicsDropShadowEffect,
)
from PyQt6.QtGui import (
    QPainter,
    QPen,
    QColor,
    QFont,
    QPixmap,
    QIcon,
)

# ── 將 py檔 資料夾加入 import 路徑 ──────────────────────────────
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

from common import get_base_path  # noqa: E402


# ── 各模式的產生函式 ─────────────────────────────────────────────
# 函式簽名：(title_text, path_templates, path_output, image_file_path, image_file_name_noext)
#   path_templates : Word 模板所在資料夾（py檔 的上一層）
#   path_output    : 產生的 .docx 要存放的資料夾（使用者選擇）


def _run_3col_landscape_filename(
    title_text, path_templates, path_output, image_file_path, image_file_name_noext
):
    import os
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_ALIGN_VERTICAL,
    )
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
    from common import (
        open_template,
        setup_header,
        set_run_font,
        open_image_as_stream,
        delete_first_paragraph_if_empty,
        delete_trailing_empty_paragraphs,
        save_document,
        fill_name_cell,
        set_cell_width,
        set_table_fixed_width,
    )

    document = open_template(os.path.join(path_templates, "word橫向模板別動.docx"))
    section = document.sections[0]
    setup_header(document, section, title_text)

    GROUP_SIZE = 3
    for group_start in range(0, len(image_file_path), GROUP_SIZE):
        group_imgs = image_file_path[group_start : group_start + GROUP_SIZE]
        group_names = image_file_name_noext[group_start : group_start + GROUP_SIZE]
        tbl = document.add_table(rows=2, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        tbl.rows[0].height = Cm(15)
        tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tbl.rows[1].height = Cm(1.2)
        tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        # 明確設定表格總寬與每個 cell 欄寬，防止 Word 自動縮放
        COL_W = 8.47
        set_table_fixed_width(tbl, COL_W * 3)
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_width(cell, COL_W)
        for i, (img_path, img_name) in enumerate(zip(group_imgs, group_names)):
            cell_pic = tbl.cell(0, i)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(img_path), width=Cm(8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_name = tbl.cell(1, i)
            fill_name_cell(
                cell_name,
                group_start + i + 1,
                f"說明：{img_name}",
                outer_width_cm=COL_W,
            )
        if group_start + GROUP_SIZE < len(image_file_path):
            document.add_section(WD_SECTION.NEW_PAGE)

    delete_first_paragraph_if_empty(document)
    delete_trailing_empty_paragraphs(document)
    return save_document(document, path_output, title_text)


def _run_3col_landscape_number(
    title_text, path_templates, path_output, image_file_path, image_file_name_noext
):
    import os
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_ALIGN_VERTICAL,
    )
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
    from common import (
        open_template,
        setup_header,
        set_run_font,
        open_image_as_stream,
        delete_first_paragraph_if_empty,
        delete_trailing_empty_paragraphs,
        save_document,
        fill_name_cell,
        set_cell_width,
        set_table_fixed_width,
    )

    document = open_template(os.path.join(path_templates, "word橫向模板別動.docx"))
    section = document.sections[0]
    setup_header(document, section, title_text)

    GROUP_SIZE = 3
    for group_start in range(0, len(image_file_path), GROUP_SIZE):
        group_imgs = image_file_path[group_start : group_start + GROUP_SIZE]
        tbl = document.add_table(rows=2, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        tbl.rows[0].height = Cm(15)
        tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tbl.rows[1].height = Cm(1.2)
        tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        # 明確設定表格總寬與每個 cell 欄寬，防止 Word 自動縮放
        COL_W = 8.47
        set_table_fixed_width(tbl, COL_W * 3)
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_width(cell, COL_W)
        for i, img_path in enumerate(group_imgs):
            cell_pic = tbl.cell(0, i)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(img_path), width=Cm(8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_name = tbl.cell(1, i)
            fill_name_cell(
                cell_name, group_start + i + 1, "說明：", outer_width_cm=COL_W
            )
        if group_start + GROUP_SIZE < len(image_file_path):
            document.add_section(WD_SECTION.NEW_PAGE)

    delete_first_paragraph_if_empty(document)
    delete_trailing_empty_paragraphs(document)
    return save_document(document, path_output, title_text)


def _run_2row_portrait_filename(
    title_text, path_templates, path_output, image_file_path, image_file_name_noext
):
    import os
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_ALIGN_VERTICAL,
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from common import (
        open_template,
        setup_header,
        set_run_font,
        open_image_as_stream,
        delete_first_paragraph_if_empty,
        delete_trailing_empty_paragraphs,
        save_document,
        fill_name_cell,
    )

    document = open_template(os.path.join(path_templates, "word模板別動.docx"))
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title_text)

    # 用寬度限制圖片（避免橫向圖超出頁面），列高由圖片自然撐開
    PIC_WIDTH = Cm(15.5)

    tbl = document.add_table(rows=len(image_file_path) * 2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, (img_path, img_name) in enumerate(
        zip(image_file_path, image_file_name_noext)
    ):
        cell_pic = tbl.cell(i * 2, 0)
        cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_pic = cell_pic.paragraphs[0]
        p_pic.add_run().add_picture(open_image_as_stream(img_path), width=PIC_WIDTH)
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_name = tbl.cell(i * 2 + 1, 0)
        fill_name_cell(cell_name, i + 1, f"說明：{img_name}", outer_width_cm=19.05)
        row_name = tbl.rows[i * 2 + 1]
        row_name.height = Cm(1.2)
        row_name.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # 圖片列設為「最小高度 10.5cm」，效果同 Word 表格內容→列→最小高度
    for i in range(len(image_file_path)):
        row_pic = tbl.rows[i * 2]
        row_pic.height = Cm(10.5)
        row_pic.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    delete_first_paragraph_if_empty(document)
    delete_trailing_empty_paragraphs(document)
    return save_document(document, path_output, title_text)


def _run_2row_portrait_number(
    title_text, path_templates, path_output, image_file_path, image_file_name_noext
):
    import os
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_ALIGN_VERTICAL,
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
    from common import (
        open_template,
        setup_header,
        set_run_font,
        open_image_as_stream,
        delete_first_paragraph_if_empty,
        delete_trailing_empty_paragraphs,
        save_document,
        fill_name_cell,
    )

    document = open_template(os.path.join(path_templates, "word模板別動.docx"))
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title_text)

    PIC_WIDTH = Cm(15.5)

    tbl = document.add_table(rows=len(image_file_path) * 2, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for i, img_path in enumerate(image_file_path):
        cell_pic = tbl.cell(i * 2, 0)
        cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_pic = cell_pic.paragraphs[0]
        p_pic.add_run().add_picture(open_image_as_stream(img_path), width=PIC_WIDTH)
        p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_name = tbl.cell(i * 2 + 1, 0)
        fill_name_cell(cell_name, i + 1, "說明：", outer_width_cm=19.05)
        row_name = tbl.rows[i * 2 + 1]
        row_name.height = Cm(1.2)
        row_name.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # 圖片列設為「最小高度 10.5cm」
    for i in range(len(image_file_path)):
        row_pic = tbl.rows[i * 2]
        row_pic.height = Cm(10.5)
        row_pic.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

    delete_first_paragraph_if_empty(document)
    delete_trailing_empty_paragraphs(document)
    return save_document(document, path_output, title_text)


def _run_2col_portrait_filename(
    title_text, path_templates, path_output, image_file_path, image_file_name_noext
):
    import os
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_ALIGN_VERTICAL,
    )
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
    from common import (
        open_template,
        setup_header,
        set_run_font,
        open_image_as_stream,
        delete_first_paragraph_if_empty,
        delete_trailing_empty_paragraphs,
        save_document,
        fill_name_cell,
        set_cell_width,
        set_table_fixed_width,
    )

    document = open_template(os.path.join(path_templates, "word模板別動.docx"))
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title_text)

    GROUP_SIZE = 2
    for group_start in range(0, len(image_file_path), GROUP_SIZE):
        group_imgs = image_file_path[group_start : group_start + GROUP_SIZE]
        group_names = image_file_name_noext[group_start : group_start + GROUP_SIZE]
        tbl = document.add_table(rows=2, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        tbl.rows[0].height = Cm(22.5)
        tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        tbl.rows[1].height = Cm(1.2)
        tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        # 明確設定表格總寬與每個 cell 欄寬，防止 Word 自動縮放
        COL_W = 9.525
        set_table_fixed_width(tbl, COL_W * 2)
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_width(cell, COL_W)
        for i, (img_path, img_name) in enumerate(zip(group_imgs, group_names)):
            cell_pic = tbl.cell(0, i)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(img_path), width=Cm(8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_name = tbl.cell(1, i)
            fill_name_cell(
                cell_name, group_start + i + 1, "說明：", outer_width_cm=COL_W
            )
        if group_start + GROUP_SIZE < len(image_file_path):
            document.add_section(WD_SECTION.NEW_PAGE)

    delete_first_paragraph_if_empty(document)
    delete_trailing_empty_paragraphs(document)
    return save_document(document, path_output, title_text)


def _run_2col_portrait_number(
    title_text, path_templates, path_output, image_file_path, image_file_name_noext
):
    import os
    from docx.enum.table import (
        WD_TABLE_ALIGNMENT,
        WD_ROW_HEIGHT_RULE,
        WD_ALIGN_VERTICAL,
    )
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm
    from common import (
        open_template,
        setup_header,
        set_run_font,
        open_image_as_stream,
        delete_first_paragraph_if_empty,
        delete_trailing_empty_paragraphs,
        save_document,
        fill_name_cell,
        set_cell_width,
        set_table_fixed_width,
    )

    document = open_template(os.path.join(path_templates, "word模板別動.docx"))
    section = document.sections[0]
    section.left_margin = section.right_margin = Cm(1.27)
    section.top_margin = section.bottom_margin = Cm(1.27)
    setup_header(document, section, title_text)

    GROUP_SIZE = 2
    for group_start in range(0, len(image_file_path), GROUP_SIZE):
        group_imgs = image_file_path[group_start : group_start + GROUP_SIZE]
        tbl = document.add_table(rows=2, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        tbl.rows[0].height = Cm(22.5)
        tbl.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        tbl.rows[1].height = Cm(1.2)
        tbl.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        # 明確設定表格總寬與每個 cell 欄寬，防止 Word 自動縮放
        COL_W = 9.525
        set_table_fixed_width(tbl, COL_W * 2)
        for row in tbl.rows:
            for cell in row.cells:
                set_cell_width(cell, COL_W)
        for i, img_path in enumerate(group_imgs):
            cell_pic = tbl.cell(0, i)
            cell_pic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p_pic = cell_pic.paragraphs[0]
            p_pic.add_run().add_picture(open_image_as_stream(img_path), width=Cm(8))
            p_pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_name = tbl.cell(1, i)
            fill_name_cell(
                cell_name, group_start + i + 1, "說明：", outer_width_cm=COL_W
            )
        if group_start + GROUP_SIZE < len(image_file_path):
            document.add_section(WD_SECTION.NEW_PAGE)

    delete_first_paragraph_if_empty(document)
    delete_trailing_empty_paragraphs(document)
    return save_document(document, path_output, title_text)


# 模式名稱 → 對應函式
MODE_MAP = {
    "1頁兩張上下（含檔名）": _run_2row_portrait_filename,
    "1頁兩張左右（含檔名）": _run_2col_portrait_filename,
    "1頁三張橫式（含檔名）": _run_3col_landscape_filename,
    "1頁兩張上下（純編號）": _run_2row_portrait_number,
    "1頁兩張左右（純編號）": _run_2col_portrait_number,
    "1頁三張橫式（純編號）": _run_3col_landscape_number,
}


# ── 排版預覽工具 ─────────────────────────────────────────────────


def _build_layout_preview(mode: str, image_paths: list, names: list):
    """
    用 Pillow 依排版模式組合模擬頁面縮圖，回傳 list[PIL.Image.Image]。
    每個元素代表一頁（RGB）。
    """
    from PIL import Image, ImageDraw, ImageFont
    import math

    # ── 頁面尺寸（像素，縮小版）──
    # 直向 A4 比例 210:297，橫向 297:210
    is_landscape = "橫式" in mode
    if is_landscape:
        PW, PH = 900, 638  # 橫向頁面
        COLS = 3
        GROUP = 3
        # 每欄可用圖片寬（PW 扣左右邊距 40*2，再扣欄間格線 2*2，除以欄數）
        COL_W = (PW - 80 - 4) // 3
        PIC_W = COL_W - 8  # 圖片寬（欄寬扣 padding）
        PIC_MAX_H = PH - 80 - 60  # 圖片最大高（扣頁邊+名稱列）
        NAME_H = 40
    elif "左右" in mode:
        PW, PH = 638, 900  # 直向
        COLS = 2
        GROUP = 2
        COL_W = (PW - 80 - 2) // 2
        PIC_W = COL_W - 8
        PIC_MAX_H = PH - 80 - 60
        NAME_H = 40
    else:  # 上下
        PW, PH = 638, 900
        COLS = 1
        GROUP = 2
        COL_W = PW - 80
        PIC_W = COL_W - 8
        PIC_MAX_H = (PH - 80 - 80) // 2  # 兩張各佔一半
        NAME_H = 40

    BG = (255, 255, 255)
    BORDER = (180, 190, 210)
    TEXT_COLOR = (50, 55, 75)
    MARGIN = 40

    def _load_thumb(path, max_w, max_h):
        """載入圖片並縮放至 max_w x max_h 內，保持比例，回傳 RGBA。"""
        try:
            img = Image.open(path).convert("RGB")
            img.thumbnail((max_w, max_h), Image.LANCZOS)
            return img
        except Exception:
            # 載入失敗：回傳灰色佔位圖
            ph = Image.new("RGB", (max_w, max_h), (200, 200, 200))
            return ph

    def _try_font(size):
        """嘗試載入系統字型，失敗時用預設字型。"""
        for name in [
            "msjh.ttc",
            "mingliu.ttc",
            "kaiu.ttf",
            "msyh.ttc",
            "arial.ttf",
        ]:
            try:
                return ImageFont.truetype(name, size)
            except Exception:
                pass
        return ImageFont.load_default()

    font_name = _try_font(14)

    pages = []
    for g_start in range(0, len(image_paths), GROUP):
        g_imgs = image_paths[g_start : g_start + GROUP]
        g_names = names[g_start : g_start + GROUP]

        page = Image.new("RGB", (PW, PH), BG)
        draw = ImageDraw.Draw(page)

        # 外框
        draw.rectangle(
            [MARGIN - 1, MARGIN - 1, PW - MARGIN, PH - MARGIN], outline=BORDER, width=1
        )

        if is_landscape or "左右" in mode:
            # ── 多欄並排 ──
            actual_cols = len(g_imgs)
            col_step = (PW - 2 * MARGIN) // COLS
            table_top = MARGIN + 10
            table_h = PH - 2 * MARGIN - 20
            pic_area_h = table_h - NAME_H

            for ci, (ipath, iname) in enumerate(zip(g_imgs, g_names)):
                cx = MARGIN + ci * col_step
                # 格子邊框
                draw.rectangle(
                    [cx, table_top, cx + col_step - 2, table_top + table_h],
                    outline=BORDER,
                    width=1,
                )
                # 圖片
                thumb = _load_thumb(ipath, col_step - 10, pic_area_h - 10)
                tw, th = thumb.size
                paste_x = cx + (col_step - tw) // 2
                paste_y = table_top + (pic_area_h - th) // 2
                page.paste(thumb, (paste_x, paste_y))
                # 名稱格
                name_y = table_top + pic_area_h
                draw.rectangle(
                    [cx, name_y, cx + col_step - 2, name_y + NAME_H],
                    outline=BORDER,
                    width=1,
                )
                draw.text(
                    (cx + col_step // 2, name_y + NAME_H // 2),
                    iname,
                    fill=TEXT_COLOR,
                    font=font_name,
                    anchor="mm",
                )
        else:
            # ── 上下兩張 ──
            row_h = (PH - 2 * MARGIN - 20) // 2
            pic_area_h = row_h - NAME_H
            col_w = PW - 2 * MARGIN

            for ri, (ipath, iname) in enumerate(zip(g_imgs, g_names)):
                ry = MARGIN + 10 + ri * row_h
                draw.rectangle(
                    [MARGIN, ry, PW - MARGIN - 1, ry + row_h - 2],
                    outline=BORDER,
                    width=1,
                )
                # 圖片
                thumb = _load_thumb(ipath, col_w - 10, pic_area_h - 10)
                tw, th = thumb.size
                paste_x = MARGIN + (col_w - tw) // 2
                paste_y = ry + (pic_area_h - th) // 2
                page.paste(thumb, (paste_x, paste_y))
                # 名稱
                name_y = ry + pic_area_h
                draw.rectangle(
                    [MARGIN, name_y, PW - MARGIN - 1, name_y + NAME_H],
                    outline=BORDER,
                    width=1,
                )
                draw.text(
                    (MARGIN + col_w // 2, name_y + NAME_H // 2),
                    iname,
                    fill=TEXT_COLOR,
                    font=font_name,
                    anchor="mm",
                )

        pages.append(page)

    return pages


class LayoutPreviewWindow(QWidget):
    """顯示排版預覽縮圖的獨立視窗（可滾動，顯示多頁）。"""

    def __init__(self, pages, parent=None):
        super().__init__(parent)
        self.setWindowTitle("排版預覽")
        self.setMinimumSize(520, 600)
        self.setStyleSheet("background-color: #2b2d3a;")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        # 提示列
        hint = QLabel(f"共 {len(pages)} 頁（模擬排版，僅供參考）")
        hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        hint.setStyleSheet("color: #a8b4d0; font-size: 11px; background: transparent;")
        layout.addWidget(hint)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea { border: none; background: #2b2d3a; }
            QScrollBar:vertical {
                background: #3a3d52; width: 8px; border-radius: 4px;
            }
            QScrollBar::handle:vertical {
                background: #6878a8; border-radius: 4px; min-height: 20px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
        """)

        container = QWidget()
        container.setStyleSheet("background: #2b2d3a;")
        vbox = QVBoxLayout(container)
        vbox.setContentsMargins(12, 12, 12, 12)
        vbox.setSpacing(16)

        for i, pil_img in enumerate(pages):
            # 轉成 QPixmap
            pil_img_rgb = pil_img.convert("RGB")
            data = pil_img_rgb.tobytes("raw", "RGB")
            from PyQt6.QtGui import QImage, QPixmap

            qimg = QImage(
                data,
                pil_img_rgb.width,
                pil_img_rgb.height,
                pil_img_rgb.width * 3,
                QImage.Format.Format_RGB888,
            )
            pixmap = QPixmap.fromImage(qimg)

            # 頁碼標籤
            page_lbl = QLabel(f"第 {i + 1} 頁")
            page_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            page_lbl.setStyleSheet(
                "color: #7888b8; font-size: 11px; background: transparent;"
            )
            vbox.addWidget(page_lbl)

            # 圖片標籤（陰影框）
            img_lbl = QLabel()
            img_lbl.setPixmap(pixmap)
            img_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
            img_lbl.setStyleSheet(
                "background: #ffffff; border: 1px solid #4a5080;"
                " border-radius: 3px; padding: 0px;"
            )
            vbox.addWidget(img_lbl)

        vbox.addStretch()
        scroll.setWidget(container)
        layout.addWidget(scroll, 1)


# ── 淺色科技感全域樣式 ────────────────────────────────────────────

STYLESHEET = """
QWidget {
    background-color: #eef1f7;
    color: #2b2d3a;
    font-family: "Segoe UI", "Microsoft JhengHei", sans-serif;
    font-size: 13px;
}

/* ── 輸入欄位（唯讀）── */
QLineEdit[readOnly="true"] {
    background-color: #f8f9fc;
    border: 1px solid #c8cfe0;
    border-radius: 3px;
    padding: 6px 10px;
    color: #2b2d3a;
}
QLineEdit[readOnly="true"]::placeholder {
    color: #5a6080;
}

/* ── 輸入欄位（可編輯）── */
QLineEdit {
    background-color: #ffffff;
    border: 1px solid #c0c7d8;
    border-radius: 3px;
    padding: 6px 10px;
    color: #2b2d3a;
    selection-background-color: #4a90d9;
    selection-color: #ffffff;
}
QLineEdit::placeholder {
    color: #5a6080;
}
QLineEdit:focus {
    border-color: #4a90d9;
    background-color: #f0f6ff;
}

/* ── 一般按鈕 ── */
QPushButton {
    background-color: #dde3f0;
    border: 1px solid #b8c0d4;
    border-radius: 4px;
    padding: 6px 14px;
    color: #3a3d52;
}
QPushButton:hover {
    background-color: #c8d0e8;
    border-color: #4a90d9;
    color: #1a5faa;
}
QPushButton:pressed {
    background-color: #4a90d9;
    border-color: #2a70b9;
    color: #ffffff;
}
QPushButton:disabled {
    background-color: #e8ebf3;
    color: #aab0c4;
    border-color: #d0d5e5;
}

/* ── 模式選擇按鈕 ── */
QPushButton#modeBtn {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #f8faff, stop:1 #dde4f2);
    border-top: 2px solid #ffffff;
    border-left: 2px solid #ffffff;
    border-right: 2px solid #8898bb;
    border-bottom: 2px solid #8898bb;
    border-radius: 5px;
    padding: 10px 4px;
    color: #1e2235;
    text-align: center;
    font-size: 12px;
}
QPushButton#modeBtn:hover {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #ffffff, stop:1 #e8eeff);
    border-top: 2px solid #ffffff;
    border-left: 2px solid #ffffff;
    border-right: 2px solid #8898bb;
    border-bottom: 2px solid #8898bb;
    color: #1e2235;
}
QPushButton#modeBtn:pressed {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #b8cce8, stop:1 #ccdaf5);
    border-top: 2px solid #6678a0;
    border-left: 2px solid #6678a0;
    border-right: 2px solid #ffffff;
    border-bottom: 2px solid #ffffff;
    color: #1e2235;
    padding-top: 12px;
    padding-left: 6px;
}
QPushButton#modeBtn:checked {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #aac8f0, stop:1 #c8deff);
    border-top: 2px solid #5580bb;
    border-left: 2px solid #5580bb;
    border-right: 2px solid #ffffff;
    border-bottom: 2px solid #ffffff;
    color: #1e2235;
    font-weight: bold;
    padding-top: 12px;
    padding-left: 6px;
}

/* ── 執行按鈕 ── */
QPushButton#runBtn {
    background-color: #2563eb;
    border: 2px solid #1d4ed8;
    border-radius: 8px;
    color: #ffffff;
    font-size: 15px;
    font-weight: bold;
    letter-spacing: 2px;
    padding: 10px;
}
QPushButton#runBtn:hover {
    background-color: #1d4ed8;
    border: 2px solid #7c3aed;
}
QPushButton#runBtn:pressed {
    background-color: #1e40af;
    border: 2px solid #1e40af;
}
QPushButton#runBtn:disabled {
    background-color: #5a6a9a;
    border: 2px solid #7a8ab8;
    color: #c8d4f0;
}

/* ── 分隔線 ── */
QFrame#separator {
    color: #c8cfe0;
}

/* ── 記錄區 ── */
QTextEdit {
    background-color: #1e2235;
    border: none;
    border-radius: 4px;
    color: #7ee787;
    font-family: "Consolas", "Courier New", monospace;
    font-size: 12px;
    padding: 6px 4px;
}

/* ── 張數標籤 ── */
QLabel#countLabel {
    color: #1a6fd4;
    font-size: 12px;
    font-weight: bold;
    background-color: #dceeff;
    border: 1px solid #4a90d9;
    border-radius: 3px;
    padding: 2px 4px;
}

/* ── 預覽清單 ── */
QListWidget#previewList {
    background-color: #e8ecf5;
    border: none;
    border-radius: 4px;
    color: #2b2d3a;
    font-size: 11px;
}
QListWidget#previewList::item {
    border-radius: 4px;
    padding: 4px 6px;
}
QListWidget#previewList::item:selected {
    background-color: #b8d0f0;
    color: #1a3a6e;
}
QListWidget#previewList::item:hover {
    background-color: #d4dff5;
}
"""


# ── HUD 科技感角框元件 ────────────────────────────────────────────


class TechFrame(QWidget):
    """
    仿 HUD 科技框：四角有 L 形缺口裝飾，標題嵌在左上角。
    內部放一個 content_widget，使用方式與 QGroupBox 相同。
    """

    # 顏色對應表（accent_key → (邊框色, 標題色, 背景色)）
    THEMES = {
        "blue": ("#4a90d9", "#1a6fd4", "#f4f7fd"),
        "teal": ("#2ab5a0", "#1a9480", "#f3fbfa"),
        "purple": ("#7c3aed", "#6d28d9", "#f7f4ff"),
        "amber": ("#d97706", "#b45309", "#fdf8ef"),
        "green": ("#16a34a", "#15803d", "#f2fbf4"),
    }

    CORNER = 6  # L 形長度
    GAP = 6  # 標題與邊框的間距
    TITLE_OFFSET_X = 15
    TITLE_OFFSET_Y = 0  # 從頂邊算

    def __init__(self, title: str, theme: str = "blue", parent=None):
        super().__init__(parent)
        border_color, title_color, bg_color = self.THEMES.get(
            theme, self.THEMES["blue"]
        )
        self._border_color = QColor(border_color)
        self._title_color = QColor(title_color)
        self._bg_color = QColor(bg_color)
        self._title = title

        # 計算標題文字寬度（用於留白）
        self._title_font = QFont("Segoe UI", 9, QFont.Weight.Bold)

        # 外層 layout（給 paintEvent 留邊距）
        outer = QVBoxLayout(self)
        outer.setContentsMargins(12, 24, 12, 10)
        outer.setSpacing(0)

        # 真正放子元件的容器
        self.content = QWidget(self)
        self.content.setStyleSheet("background: transparent;")
        outer.addWidget(self.content)

        # 預留一個內部 layout 給使用方填入
        self._inner = QHBoxLayout(self.content)
        self._inner.setContentsMargins(0, 0, 0, 0)
        self._inner.setSpacing(6)

    def inner_layout(self):
        """回傳內部 QHBoxLayout，供外部 addWidget 使用。"""
        return self._inner

    def paintEvent(self, event):
        super().paintEvent(event)
        p = QPainter(self)
        p.setRenderHint(QPainter.RenderHint.Antialiasing, False)

        w, h = self.width(), self.height()
        c = self.CORNER

        # ── 先量文字寬度 ──
        p.setFont(self._title_font)
        text_w = p.fontMetrics().horizontalAdvance(self._title)
        title_fm_width = text_w + 20
        title_start = self.TITLE_OFFSET_X
        title_end = title_start + title_fm_width

        # ── 先畫標題文字 ──
        BORDER_Y = 10  # 框線的 y 座標
        TITLE_H = 18  # 文字區塊高度
        title_y = BORDER_Y - TITLE_H // 2  # 文字垂直置中於框線

        p.setPen(QPen(self._title_color))
        p.drawText(
            title_start,
            title_y,
            title_fm_width,
            TITLE_H,
            Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignHCenter,
            self._title,
        )

        # ── 再畫框線 ──
        pen = QPen(self._border_color, 1.5)
        p.setPen(pen)

        p.drawLine(c, BORDER_Y, title_start - 4, BORDER_Y)  # 上邊左段
        p.drawLine(title_end + 4, BORDER_Y, w - c, BORDER_Y)  # 上邊右段
        p.drawLine(1, BORDER_Y, 1, h - c)  # 左邊
        p.drawLine(w - 1, BORDER_Y, w - 1, h - c)  # 右邊
        p.drawLine(c, h - 1, w - c, h - 1)  # 下邊

        # ── 四角 L 形 ──
        # 左上
        p.drawLine(1, BORDER_Y, 1, BORDER_Y + c)
        p.drawLine(1, BORDER_Y, c, BORDER_Y)
        # 右上
        p.drawLine(w - c, BORDER_Y, w - 1, BORDER_Y)
        p.drawLine(w - 1, BORDER_Y, w - 1, BORDER_Y + c)
        # 左下
        p.drawLine(1, h - c, 1, h - 1)
        p.drawLine(1, h - 1, c, h - 1)
        # 右下
        p.drawLine(w - c, h - 1, w - 1, h - 1)
        p.drawLine(w - 1, h - c, w - 1, h - 1)

        # ── 右下角小方塊裝飾 ──
        pen2 = QPen(self._border_color, 1)
        p.setPen(pen2)
        p.drawRect(w - 8, h - 8, 5, 5)

        p.end()


# ── 可拖曳排序的圖片卡片系統 ─────────────────────────────────────

from PyQt6.QtCore import pyqtSignal as _Signal, QMimeData, QByteArray, QPoint
from PyQt6.QtGui import QDrag, QCursor
from PyQt6.QtWidgets import QScrollArea, QScrollBar


class PhotoCard(QWidget):
    """
    單張圖片卡片：編號徽章 + 縮圖 + 可雙擊編輯的檔名。
    拖曳時攜帶自身在 container 中的 index。
    """

    name_changed = _Signal(int, str)  # (index, new_name)
    remove_requested = _Signal(int)  # (index)

    def __init__(self, path: str, name: str, index: int, parent=None):
        super().__init__(parent)
        self.path = path
        self.name = name  # 顯示名稱（可被使用者編輯）
        self.index = index
        self._drag_start: QPoint | None = None
        self._dragging: bool = False
        self._hovered: bool = False

        self.setFixedWidth(280)
        self.setCursor(QCursor(Qt.CursorShape.OpenHandCursor))

        outer = QVBoxLayout(self)
        outer.setContentsMargins(6, 6, 6, 6)
        outer.setSpacing(4)

        # ── 縮圖容器（用來疊放編號徽章）────────────────────────
        img_container = QWidget()
        img_container.setFixedHeight(200)
        img_container.setStyleSheet("background: transparent;")
        img_stack = QVBoxLayout(img_container)
        img_stack.setContentsMargins(0, 0, 0, 0)

        self.img_label = QLabel()
        self.img_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.img_label.setFixedHeight(200)
        self.img_label.setStyleSheet("background: #d8dff0; border-radius: 4px;")
        px = QPixmap(path)
        if not px.isNull():
            px = px.scaled(
                268,
                196,
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation,
            )
        else:
            px = QPixmap(268, 196)
            px.fill(QColor("#c0c8e0"))
        self.img_label.setPixmap(px)
        img_stack.addWidget(self.img_label)

        # 編號徽章（絕對定位，疊在縮圖左上角）
        self.badge = QLabel(str(index + 1), img_container)
        self.badge.setFixedSize(26, 26)
        self.badge.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.badge.move(6, 6)
        self.badge.setStyleSheet("""
            background-color: #0891b2;
            color: #ffffff;
            font-size: 11px;
            font-weight: bold;
            border-radius: 13px;
        """)
        self.badge.raise_()

        # ✕ 刪除按鈕（絕對定位，疊在縮圖右上角）
        self.btn_remove = QPushButton("✕", img_container)
        self.btn_remove.setFixedSize(22, 22)
        self.btn_remove.move(240, 6)
        self.btn_remove.setStyleSheet("""
            QPushButton {
                background-color: #e53e3e;
                color: #ffffff;
                font-size: 10px;
                font-weight: bold;
                border-radius: 11px;
                border: none;
                padding: 0px;
            }
            QPushButton:hover { background-color: #c53030; }
            QPushButton:pressed { background-color: #9b2c2c; }
        """)
        self.btn_remove.setCursor(QCursor(Qt.CursorShape.PointingHandCursor))
        self.btn_remove.clicked.connect(lambda: self.remove_requested.emit(self.index))
        self.btn_remove.raise_()

        outer.addWidget(img_container)

        # ── 名稱列（藍框容器：編號 + 名稱 label / 編輯框）──────
        name_box = QFrame()
        name_box.setStyleSheet(
            "QFrame { background: #ffffff; border: 1px solid #4a90d9;"
            " border-radius: 4px; }"
        )
        name_box.setFixedHeight(26)
        name_row = QHBoxLayout(name_box)
        name_row.setContentsMargins(4, 0, 4, 0)
        name_row.setSpacing(3)

        self.num_label = QLabel(f"{index + 1}.")
        self.num_label.setFixedWidth(22)
        self.num_label.setAlignment(
            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
        )
        self.num_label.setStyleSheet(
            "color: #0891b2; font-size: 11px; font-weight: bold;"
            " background: transparent; border: none;"
        )
        name_row.addWidget(self.num_label)

        self.name_label = QLabel(name)
        self.name_label.setAlignment(
            Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter
        )
        self.name_label.setWordWrap(False)
        self.name_label.setStyleSheet(
            "color: #2b2d3a; font-size: 11px; background: transparent; border: none;"
        )
        self.name_label.setToolTip("雙擊可編輯名稱")
        name_row.addWidget(self.name_label, 1)

        self.name_edit = QLineEdit(name)
        self.name_edit.setStyleSheet(
            "color: #2b2d3a; font-size: 11px; background: transparent;"
            " border: none; padding: 0px;"
        )
        self.name_edit.hide()
        self.name_edit.returnPressed.connect(self._commit_edit)
        self.name_edit.editingFinished.connect(self._commit_edit)
        name_row.addWidget(self.name_edit, 1)

        outer.addWidget(name_box)

        self.setStyleSheet("""
            PhotoCard {
                background: #eef1f7;
                border: 1px solid #c8d0e8;
                border-radius: 6px;
            }
            PhotoCard:hover {
                border: 1px solid #4a90d9;
                background: #e4ecfa;
            }
        """)

        # 陰影效果（預設輕微）
        self._shadow = QGraphicsDropShadowEffect(self)
        self._shadow.setBlurRadius(8)
        self._shadow.setOffset(0, 2)
        self._shadow.setColor(QColor(0, 0, 0, 40))
        self.setGraphicsEffect(self._shadow)

        # 位移動畫
        self._anim = QPropertyAnimation(self, b"pos")
        self._anim.setDuration(120)
        self._anim.setEasingCurve(QEasingCurve.Type.OutCubic)
        self._base_pos: QPoint | None = None

    def set_index(self, index: int):
        """更新顯示的編號（重排後呼叫）。"""
        self.index = index
        self.badge.setText(str(index + 1))
        self.num_label.setText(f"{index + 1}.")

    # ── 雙擊編輯名稱 ──────────────────────────────────────────────

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._start_edit()
        super().mouseDoubleClickEvent(event)

    def _start_edit(self):
        self.name_edit.setText(self.name)
        self.name_label.hide()
        self.name_edit.show()
        self.name_edit.selectAll()
        self.name_edit.setFocus()
        self.setCursor(QCursor(Qt.CursorShape.IBeamCursor))

    def _commit_edit(self):
        new_name = self.name_edit.text().strip()
        if not new_name:
            new_name = self.name  # 空白時還原
        self.name = new_name
        self.name_label.setText(new_name)
        self.name_edit.hide()
        self.name_label.show()
        self.setCursor(QCursor(Qt.CursorShape.OpenHandCursor))
        self.name_changed.emit(self.index, new_name)

    # ── Hover 浮起效果 ────────────────────────────────────────────

    def _set_hover(self, on: bool):
        """切換 hover 浮起狀態：陰影加深 + 卡片上移 4px。"""
        self._hovered = on
        self._anim.stop()
        if on:
            self._base_pos = self.pos()  # 記錄 layout 分配的原始位置
            self._rebuild_shadow(blur=20, dy=6, alpha=80)
            target = self._base_pos + QPoint(0, -4)
            self._anim.setStartValue(self.pos())
            self._anim.setEndValue(target)
            self._anim.start()
        else:
            self._rebuild_shadow(blur=8, dy=2, alpha=40)
            if self._base_pos is not None:
                self._anim.setStartValue(self.pos())
                self._anim.setEndValue(self._base_pos)
                self._anim.start()
                # 動畫結束後強制歸位，防止浮點誤差殘留
                self._anim.finished.connect(self._snap_to_base)

    def _rebuild_shadow(self, blur: int, dy: int, alpha: int):
        """重新建立 shadow effect（避免舊物件被 Qt 刪除後存取崩潰）。"""
        self._shadow = QGraphicsDropShadowEffect(self)
        self._shadow.setBlurRadius(blur)
        self._shadow.setOffset(0, dy)
        self._shadow.setColor(QColor(0, 0, 0, alpha))
        self.setGraphicsEffect(self._shadow)

    def _snap_to_base(self):
        """動畫結束後把 pos 強制歸回 layout 的原始位置。"""
        try:
            self._anim.finished.disconnect(self._snap_to_base)
        except Exception:
            pass
        if self._base_pos is not None and not self._hovered:
            self.move(self._base_pos)
            self._base_pos = None

    def enterEvent(self, event):
        if not self._dragging:
            self._set_hover(True)
        super().enterEvent(event)

    def leaveEvent(self, event):
        if self._hovered:
            self._set_hover(False)
        super().leaveEvent(event)

    # ── 拖曳邏輯 ──────────────────────────────────────────────────

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._drag_start = event.pos()
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if (
            not self._dragging
            and self._drag_start is not None
            and event.buttons() & Qt.MouseButton.LeftButton
        ):
            dist = (event.pos() - self._drag_start).manhattanLength()
            if dist >= 8:
                self._start_drag()
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        self._drag_start = None
        super().mouseReleaseEvent(event)

    def _start_drag(self):
        # 先強制還原 hover 位移，再啟動拖曳
        if self._hovered or self._base_pos is not None:
            self._anim.stop()
            self._hovered = False
            if self._base_pos is not None:
                self.move(self._base_pos)
                self._base_pos = None
        self._dragging = True
        # 被拖卡片變半透明（暫時覆蓋 shadow effect）
        opacity = QGraphicsOpacityEffect(self)
        opacity.setOpacity(0.4)
        self.setGraphicsEffect(opacity)
        self.setCursor(QCursor(Qt.CursorShape.ClosedHandCursor))
        drag = QDrag(self)
        mime = QMimeData()
        mime.setData(
            "application/x-photocard-index", QByteArray(str(self.index).encode())
        )
        drag.setMimeData(mime)
        # 拖曳預覽圖：縮小版縮圖
        preview = self.img_label.pixmap()
        if preview and not preview.isNull():
            drag.setPixmap(
                preview.scaled(
                    100,
                    80,
                    Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation,
                )
            )
            drag.setHotSpot(QPoint(50, 40))
        drag.exec(Qt.DropAction.MoveAction)
        # 拖曳結束：重新建立陰影（舊的被 Qt 在 setGraphicsEffect(opacity) 時自動刪除）
        self._rebuild_shadow(blur=8, dy=2, alpha=40)
        self._dragging = False
        self._drag_start = None
        self.setCursor(QCursor(Qt.CursorShape.OpenHandCursor))


class CardContainer(QWidget):
    """
    垂直排列的卡片容器，接受 PhotoCard 拖曳後即時重排（iOS 風格）。
    排序完成後發出 order_changed(list[str]) 訊號，攜帶新的 path 順序。
    """

    order_changed = _Signal(list)
    name_changed = _Signal(int, str)  # (index, new_name)
    card_removed = _Signal(int)  # (index) 某張被刪除後發出

    def __init__(self, parent=None):
        super().__init__(parent)
        self._cards: list[PhotoCard] = []
        self._drag_src: int = -1  # 正在被拖的卡片原始 index
        self._last_swap: int = -1  # 上次換到的目標 index（debounce 用）

        self.setAcceptDrops(True)
        self._layout = QVBoxLayout(self)
        self._layout.setContentsMargins(4, 4, 4, 4)
        self._layout.setSpacing(8)
        self._layout.addStretch()

    # ── 公開 API ──────────────────────────────────────────────────

    def set_images(self, paths: list[str], names: list[str]):
        """清空並重新建立所有卡片。"""
        for card in self._cards:
            self._layout.removeWidget(card)
            card.deleteLater()
        self._cards = []

        for i, (path, name) in enumerate(zip(paths, names)):
            card = PhotoCard(path, name, i)
            card.name_changed.connect(self._on_card_name_changed)
            card.remove_requested.connect(self.remove_card)
            self._layout.insertWidget(i, card)
            self._cards.append(card)

        self.setMinimumHeight(len(self._cards) * (220 + 8) + 8)

    def paths(self) -> list[str]:
        return [c.path for c in self._cards]

    def names(self) -> list[str]:
        return [c.name for c in self._cards]

    def remove_card(self, index: int):
        """移除指定 index 的卡片，重排剩餘卡片編號，發出 card_removed 訊號。"""
        if index < 0 or index >= len(self._cards):
            return
        card = self._cards.pop(index)
        self._layout.removeWidget(card)
        card.deleteLater()
        # 更新剩餘卡片編號
        for i, c in enumerate(self._cards):
            c.index = i
            c.set_index(i)
        self.setMinimumHeight(len(self._cards) * (220 + 8) + 8)
        self.card_removed.emit(index)

    def _on_card_name_changed(self, idx: int, new_name: str):
        """轉發卡片的 name_changed 訊號給外部（MainWindow）。"""
        self.name_changed.emit(idx, new_name)

    # ── 位置計算 ──────────────────────────────────────────────────

    def _card_at(self, y: int) -> int:
        """回傳 y 座標對應的卡片 index（用 layout geometry，不受 hover 動畫影響）。"""
        for i in range(len(self._cards)):
            item = self._layout.itemAt(i)
            if item is None:
                continue
            geo = item.geometry()
            if geo.top() <= y <= geo.bottom():
                return i
        return -1

    # ── 即時重排 ──────────────────────────────────────────────────

    def _reorder(self, src: int, dst: int):
        """把 src 卡片移到 dst 位置，即時更新 layout。"""
        if src == dst or src < 0 or dst < 0:
            return
        card = self._cards.pop(src)
        self._cards.insert(dst, card)
        for c in self._cards:
            self._layout.removeWidget(c)
        for i, c in enumerate(self._cards):
            self._layout.insertWidget(i, c)
            c.index = i
            c.set_index(i)

    # ── Drop 事件 ─────────────────────────────────────────────────

    def dragEnterEvent(self, event):
        if not event.mimeData().hasFormat("application/x-photocard-index"):
            event.ignore()
            return
        self._drag_src = int(
            event.mimeData().data("application/x-photocard-index").data().decode()
        )
        self._last_swap = self._drag_src
        event.acceptProposedAction()

    def dragMoveEvent(self, event):
        if not event.mimeData().hasFormat("application/x-photocard-index"):
            event.ignore()
            return
        y = event.position().toPoint().y()
        dst = self._card_at(y)
        # 只有游標移到不同卡片上才觸發重排，避免抖動
        if dst != -1 and dst != self._last_swap:
            self._reorder(self._last_swap, dst)
            self._last_swap = dst
        event.acceptProposedAction()

    def dragLeaveEvent(self, event):
        # 拖曳離開容器：把卡片還原到拖曳前的位置
        if self._drag_src != -1 and self._last_swap != self._drag_src:
            self._reorder(self._last_swap, self._drag_src)
            self._last_swap = self._drag_src

    def dropEvent(self, event):
        if not event.mimeData().hasFormat("application/x-photocard-index"):
            event.ignore()
            return
        # 最終位置已在 dragMoveEvent 即時更新，直接確認
        self._drag_src = -1
        self._last_swap = -1
        event.acceptProposedAction()
        self.order_changed.emit(self.paths())


# 模式按鈕顯示文字（順序對應 MODE_MAP）
MODE_LABELS = [
    "⊟  兩張上下\n含檔名",
    "⊠  兩張左右\n含檔名",
    "⊞  三張橫式\n含檔名",
    "⊟  兩張上下\n純編號",
    "⊠  兩張左右\n純編號",
    "⊞  三張橫式\n純編號",
]


# ── 背景執行緒（避免 UI 凍結）────────────────────────────────────


class WorkerThread(QThread):
    log = pyqtSignal(str)
    success = pyqtSignal(str)
    error = pyqtSignal(str)

    def __init__(
        self,
        mode,
        title,
        path_templates,
        path_output,
        image_file_path,
        image_file_name_noext,
    ):
        super().__init__()
        self.mode = mode
        self.title = title
        self.path_templates = path_templates
        self.path_output = path_output
        self.image_file_path = image_file_path
        self.image_file_name_noext = image_file_name_noext

    def run(self):
        try:
            self.log.emit("開始產生文件...")
            func = MODE_MAP[self.mode]
            output = func(
                self.title,
                self.path_templates,
                self.path_output,
                self.image_file_path,
                self.image_file_name_noext,
            )
            self.success.emit(output)
        except Exception:
            self.error.emit(traceback.format_exc())


# ── 主視窗 ───────────────────────────────────────────────────────


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self._here = os.path.dirname(os.path.abspath(__file__))
        self.path_templates = get_base_path()
        self.image_file_path = []
        self.image_file_name_noext = []
        self.custom_names: list[str] = []
        self.path_output = ""
        self._worker = None
        self._selected_mode = list(MODE_MAP.keys())[0]
        self._build_ui()

    # ── UI 建構 ──────────────────────────────────────────────────

    def _build_ui(self):
        self.setWindowTitle("小工具")
        self.setMinimumWidth(560)
        self.setStyleSheet(STYLESHEET)
        self._preview_expanded = False  # 記錄預覽面板是否已展開

        root = QVBoxLayout(self)
        root.setSpacing(10)
        root.setContentsMargins(20, 14, 20, 14)

        # ── 頂部標題列 ──
        title_bar = QHBoxLayout()
        app_title = QLabel("📷 酷酷的照片黏貼表")
        app_title.setStyleSheet(
            "color: #1a3a6e; font-size: 15px; font-weight: bold; letter-spacing: 2px;"
            "background: transparent;"
        )
        subtitle = QLabel("v2.0")
        subtitle.setStyleSheet(
            "color: #4a90d9; font-size: 10px; font-weight: bold;"
            "background: #dceeff; border: 1px solid #4a90d9;"
            "border-radius: 3px; padding: 1px 5px;"
        )
        credit = QLabel("ft.林瑾孝")
        credit.setStyleSheet(
            "color: #8898bb; font-size: 12px; background: transparent;"
        )
        title_bar.addWidget(app_title)
        title_bar.addSpacing(8)
        title_bar.addWidget(subtitle)
        title_bar.addStretch()
        title_bar.addWidget(credit)
        root.addLayout(title_bar)

        sep = QFrame()
        sep.setObjectName("separator")
        sep.setFrameShape(QFrame.Shape.HLine)
        root.addWidget(sep)

        # ── 主體水平佈局（左側操作區 + 右側預覽面板）──
        body_layout = QHBoxLayout()
        body_layout.setSpacing(12)
        body_layout.setContentsMargins(0, 0, 0, 0)

        # ── 左側主操作區 ──
        left_widget = QWidget()
        left_widget.setStyleSheet("background: transparent;")
        left = QVBoxLayout(left_widget)
        left.setSpacing(10)
        left.setContentsMargins(0, 0, 0, 0)

        # 照片來源（TechFrame blue）
        folder_frame = TechFrame("照片來源", theme="blue")
        fl = folder_frame.inner_layout()
        fl.setSpacing(8)
        self.txt_folder = QLineEdit()
        self.txt_folder.setReadOnly(True)
        self.txt_folder.setPlaceholderText("尚未選擇圖片...")
        self.lbl_count = QLabel("")
        self.lbl_count.setObjectName("countLabel")
        self.lbl_count.setFixedWidth(52)
        self.lbl_count.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.lbl_count.setVisible(False)
        self.btn_browse_input = QPushButton("選擇圖片")
        self.btn_browse_input.setFixedWidth(88)
        self.btn_browse_input.clicked.connect(self._on_browse_input)
        self.btn_clear_input = QPushButton("清除")
        self.btn_clear_input.setFixedWidth(52)
        self.btn_clear_input.setVisible(False)
        self.btn_clear_input.setStyleSheet(
            "QPushButton { background-color: #e53e3e; border: 1px solid #c53030;"
            " border-radius: 4px; color: #ffffff; font-size: 12px; padding: 6px 8px; }"
            "QPushButton:hover { background-color: #c53030; }"
            "QPushButton:pressed { background-color: #9b2c2c; }"
        )
        self.btn_clear_input.clicked.connect(self._on_clear_input)
        fl.addWidget(self.txt_folder, 1)
        fl.addWidget(self.lbl_count)
        fl.addWidget(self.btn_clear_input)
        fl.addWidget(self.btn_browse_input)
        left.addWidget(folder_frame)

        # 輸出資料夾（TechFrame teal）
        output_frame = TechFrame("輸出資料夾", theme="teal")
        ol = output_frame.inner_layout()
        ol.setSpacing(8)
        self.txt_output = QLineEdit()
        self.txt_output.setReadOnly(True)
        self.txt_output.setPlaceholderText("尚未選擇資料夾...")
        self.btn_browse_output = QPushButton("選擇資料夾")
        self.btn_browse_output.setFixedWidth(100)
        self.btn_browse_output.clicked.connect(self._on_browse_output)
        ol.addWidget(self.txt_output, 1)
        ol.addWidget(self.btn_browse_output)
        left.addWidget(output_frame)

        # 頁首標題（TechFrame purple）
        title_frame = TechFrame("頁首標題", theme="purple")
        tl = title_frame.inner_layout()
        self.input_title = QLineEdit()
        self.input_title.setPlaceholderText("輸入要顯示在頁首的標題文字...")
        tl.addWidget(self.input_title)
        left.addWidget(title_frame)

        # 排版模式（TechFrame amber，3×2 格）
        mode_frame = TechFrame("排版模式", theme="amber")
        mode_outer = mode_frame.inner_layout()
        mode_grid_widget = QWidget()
        mode_grid_widget.setStyleSheet("background: transparent;")
        mode_grid = QGridLayout(mode_grid_widget)
        mode_grid.setSpacing(8)
        mode_grid.setContentsMargins(0, 0, 0, 0)

        self._mode_btn_group = QButtonGroup(self)
        self._mode_btn_group.setExclusive(True)
        mode_keys = list(MODE_MAP.keys())

        for idx, (key, label) in enumerate(zip(mode_keys, MODE_LABELS)):
            btn = QPushButton(label)
            btn.setObjectName("modeBtn")
            btn.setCheckable(True)
            btn.setFixedHeight(58)
            btn.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
            if idx == 0:
                btn.setChecked(True)
            self._mode_btn_group.addButton(btn, idx)
            mode_grid.addWidget(btn, idx // 3, idx % 3)

        self._mode_btn_group.idClicked.connect(self._on_mode_selected)
        mode_outer.addWidget(mode_grid_widget)
        left.addWidget(mode_frame)

        # 預覽排版按鈕
        self.btn_preview_layout = QPushButton("🔍  預覽排版")
        self.btn_preview_layout.setFixedHeight(36)
        self.btn_preview_layout.setStyleSheet("""
            QPushButton {
                background-color: #3a4a6e;
                border: 1px solid #4a5a8e;
                border-radius: 6px;
                color: #a8c0f0;
                font-size: 13px;
                font-weight: bold;
                letter-spacing: 1px;
                padding: 6px;
            }
            QPushButton:hover {
                background-color: #4a5a8e;
                border-color: #6878b8;
                color: #d0e0ff;
            }
            QPushButton:pressed {
                background-color: #5a6a9e;
            }
            QPushButton:disabled {
                background-color: #2e3048;
                color: #5a6080;
                border-color: #3a4060;
            }
        """)
        self.btn_preview_layout.clicked.connect(self._on_preview_layout)
        left.addWidget(self.btn_preview_layout)

        # 執行按鈕
        self.btn_run = QPushButton("▶  產生 Word 文件")
        self.btn_run.setObjectName("runBtn")
        self.btn_run.setFixedHeight(48)
        self.btn_run.setStyleSheet("""
            QPushButton {
                background-color: #0891b2;
                border: 2px solid #0e7490;
                border-radius: 8px;
                color: #ffffff;
                font-size: 15px;
                font-weight: bold;
                letter-spacing: 2px;
                padding: 10px;
            }
            QPushButton:hover {
                background-color: #0e7490;
                border: 2px solid #164e63;
            }
            QPushButton:pressed {
                background-color: #164e63;
                border: 2px solid #164e63;
            }
            QPushButton:disabled {
                background-color: #5a6a9a;
                border: 2px solid #7a8ab8;
                color: #c8d4f0;
            }
        """)
        self.btn_run.clicked.connect(self._on_run)
        left.addWidget(self.btn_run)

        # 記錄區（TechFrame green）
        log_frame = TechFrame("執行記錄", theme="green")
        ll = log_frame.inner_layout()
        self.log_box = QTextEdit()
        self.log_box.setReadOnly(True)
        self.log_box.setFixedHeight(90)
        self.log_box.document().setDocumentMargin(0)
        self.log_box.setHtml('<span style="color:#7ee787;">&gt; 等待執行...</span>')
        ll.addWidget(self.log_box)
        left.addWidget(log_frame)

        body_layout.addWidget(left_widget)

        # ── 右側預覽面板（預設隱藏，選圖後展開）──
        self.preview_frame = TechFrame("圖片預覽（點兩下可編輯檔名）", theme="blue")
        self.preview_frame.setFixedWidth(320)
        self.preview_frame.setVisible(False)

        preview_inner = self.preview_frame.inner_layout()
        preview_inner.setContentsMargins(0, 0, 0, 0)

        # 整個右側面板用垂直 layout 包住（工具列 + 捲動區）
        preview_vbox_widget = QWidget()
        preview_vbox_widget.setStyleSheet("background: transparent;")
        preview_vbox = QVBoxLayout(preview_vbox_widget)
        preview_vbox.setContentsMargins(0, 0, 0, 0)
        preview_vbox.setSpacing(6)

        # ── 批量改名工具列 ──────────────────────────────────────
        rename_bar = QFrame()
        rename_bar.setStyleSheet(
            "QFrame { background: #dce8f8; border: 1px solid #4a90d9;"
            " border-radius: 5px; }"
        )
        rename_bar_layout = QVBoxLayout(rename_bar)
        rename_bar_layout.setContentsMargins(8, 6, 8, 6)
        rename_bar_layout.setSpacing(5)

        # 標題列
        rename_title = QLabel("批量改名")
        rename_title.setStyleSheet(
            "color: #1a5faa; font-size: 11px; font-weight: bold;"
            " background: transparent; border: none;"
        )
        rename_bar_layout.addWidget(rename_title)

        # 新名稱輸入列
        name_input_row = QHBoxLayout()
        name_input_row.setSpacing(4)
        lbl_new = QLabel("新名稱：")
        lbl_new.setStyleSheet(
            "color: #2b2d3a; font-size: 11px; background: transparent; border: none;"
        )
        lbl_new.setFixedWidth(46)
        self.rename_new_name = QLineEdit()
        self.rename_new_name.setPlaceholderText("例：照片")
        self.rename_new_name.setFixedHeight(24)
        self.rename_new_name.setStyleSheet(
            "background: #ffffff; border: 1px solid #4a90d9;"
            " border-radius: 3px; padding: 1px 5px; font-size: 11px;"
        )
        name_input_row.addWidget(lbl_new)
        name_input_row.addWidget(self.rename_new_name, 1)
        rename_bar_layout.addLayout(name_input_row)

        # 範圍輸入列
        range_row = QHBoxLayout()
        range_row.setSpacing(4)
        lbl_from = QLabel("第")
        lbl_from.setStyleSheet(
            "color: #2b2d3a; font-size: 11px; background: transparent; border: none;"
        )
        self.rename_from = QLineEdit("1")
        self.rename_from.setFixedWidth(36)
        self.rename_from.setFixedHeight(24)
        self.rename_from.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.rename_from.setStyleSheet(
            "background: #ffffff; border: 1px solid #4a90d9;"
            " border-radius: 3px; padding: 1px 3px; font-size: 11px;"
        )
        lbl_to = QLabel("張 到 第")
        lbl_to.setStyleSheet(
            "color: #2b2d3a; font-size: 11px; background: transparent; border: none;"
        )
        self.rename_to = QLineEdit("1")
        self.rename_to.setFixedWidth(36)
        self.rename_to.setFixedHeight(24)
        self.rename_to.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.rename_to.setStyleSheet(
            "background: #ffffff; border: 1px solid #4a90d9;"
            " border-radius: 3px; padding: 1px 3px; font-size: 11px;"
        )
        lbl_end = QLabel("張")
        lbl_end.setStyleSheet(
            "color: #2b2d3a; font-size: 11px; background: transparent; border: none;"
        )
        btn_apply_rename = QPushButton("套用")
        btn_apply_rename.setFixedHeight(24)
        btn_apply_rename.setFixedWidth(44)
        btn_apply_rename.setStyleSheet(
            "QPushButton { background: #0891b2; border: 1px solid #0e7490;"
            " border-radius: 3px; color: #ffffff; font-size: 11px;"
            " font-weight: bold; padding: 0px; }"
            "QPushButton:hover { background: #0e7490; }"
            "QPushButton:pressed { background: #164e63; }"
        )
        btn_apply_rename.clicked.connect(self._on_batch_rename)
        range_row.addWidget(lbl_from)
        range_row.addWidget(self.rename_from)
        range_row.addWidget(lbl_to)
        range_row.addWidget(self.rename_to)
        range_row.addWidget(lbl_end)
        range_row.addStretch()
        range_row.addWidget(btn_apply_rename)
        rename_bar_layout.addLayout(range_row)

        preview_vbox.addWidget(rename_bar)

        # ── CardContainer 放在 QScrollArea 內，可縱向捲動 ─────
        self.card_container = CardContainer()
        self.card_container.order_changed.connect(self._on_preview_reordered)
        self.card_container.name_changed.connect(self._on_card_name_changed)
        self.card_container.card_removed.connect(self._on_card_removed)

        scroll = QScrollArea()
        scroll.setWidget(self.card_container)
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        scroll.setStyleSheet("""
            QScrollArea { border: none; background: transparent; }
            QScrollArea > QWidget > QWidget { background: transparent; }
            QScrollBar:vertical {
                background: #d8dff0; width: 8px; border-radius: 4px;
            }
            QScrollBar::handle:vertical {
                background: #8898cc; border-radius: 4px; min-height: 20px;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
        """)
        preview_vbox.addWidget(scroll, 1)

        preview_inner.addWidget(preview_vbox_widget)

        body_layout.addWidget(self.preview_frame)

        root.addLayout(body_layout)

    # ── 事件處理 ─────────────────────────────────────────────────

    def _on_preview_layout(self):
        """依目前排版模式和已選圖片產生預覽視窗。"""
        if not self.image_file_path:
            QMessageBox.warning(self, "未選擇圖片", "請先選擇照片再預覽。")
            return
        try:
            pages = _build_layout_preview(
                self._selected_mode,
                self.image_file_path,
                self.custom_names if self.custom_names else self.image_file_name_noext,
            )
        except Exception as e:
            QMessageBox.critical(self, "預覽失敗", str(e))
            return
        self._preview_win = LayoutPreviewWindow(pages, parent=None)
        self._preview_win.show()

    def _on_mode_selected(self, idx):
        self._selected_mode = list(MODE_MAP.keys())[idx]

    def _on_browse_input(self):
        from common import VALID_IMAGE_EXTS

        exts = " ".join(f"*{e}" for e in sorted(VALID_IMAGE_EXTS, key=str.lower))
        start_dir = (
            os.path.dirname(self.image_file_path[0])
            if self.image_file_path
            else self.path_templates
        )
        files, _ = QFileDialog.getOpenFileNames(
            self,
            "選擇圖片（可多選）",
            start_dir,
            f"圖片檔案 ({exts})",
        )
        if not files:
            return
        files.sort()
        self.image_file_path = files
        self.image_file_name_noext = [
            os.path.splitext(os.path.basename(f))[0] for f in files
        ]
        count = len(files)
        self.lbl_count.setText(f"{count} 張")
        self.lbl_count.setVisible(True)
        self.btn_clear_input.setVisible(True)
        if count == 1:
            self.txt_folder.setText(files[0])
        else:
            self.txt_folder.setText(os.path.dirname(files[0]))

        # ── 預設輸出資料夾 = 圖片所在資料夾（使用者未手動選過才覆蓋）──
        default_output = os.path.dirname(files[0])
        if not self.path_output:
            self.path_output = default_output
            self.txt_output.setText(default_output)

        # ── 填入預覽縮圖 ──
        self._refresh_preview()
        # 批量改名範圍預設為全部
        self.rename_from.setText("1")
        self.rename_to.setText(str(len(self.image_file_path)))

        if not self._preview_expanded:
            # 第一次選圖：顯示面板並向右展開視窗
            self.preview_frame.setVisible(True)
            self.resize(self.width() + 332, self.height())
            self._preview_expanded = True
        else:
            # 已展開：只更新縮圖，不改視窗大小
            pass

    def _refresh_preview(self):
        """重新從 image_file_path 建立卡片清單。"""
        self.custom_names = list(self.image_file_name_noext)
        self.card_container.set_images(self.image_file_path, self.custom_names)

    def _on_preview_reordered(self, new_paths: list):
        """拖曳排序後，同步更新 image_file_path / image_file_name_noext / custom_names。"""
        # 建立舊 path → custom_name 的對照，保留使用者已編輯的名稱
        old_name_map = {c.path: c.name for c in self.card_container._cards}
        self.image_file_path = new_paths
        self.image_file_name_noext = [
            os.path.splitext(os.path.basename(p))[0] for p in new_paths
        ]
        self.custom_names = [
            old_name_map[p]
            if p in old_name_map
            else os.path.splitext(os.path.basename(p))[0]
            for p in new_paths
        ]
        self.lbl_count.setText(f"{len(new_paths)} 張")

    def _on_card_name_changed(self, idx: int, new_name: str):
        """卡片名稱被使用者編輯後，更新 custom_names。"""
        if 0 <= idx < len(self.custom_names):
            self.custom_names[idx] = new_name

    def _on_card_removed(self, idx: int):
        """個別卡片被 ✕ 刪除後，同步 image_file_path / custom_names。"""
        if 0 <= idx < len(self.image_file_path):
            self.image_file_path.pop(idx)
            self.image_file_name_noext.pop(idx)
        if 0 <= idx < len(self.custom_names):
            self.custom_names.pop(idx)
        count = len(self.image_file_path)
        if count == 0:
            self._reset_input_state()
        else:
            self.lbl_count.setText(f"{count} 張")
            if count == 1:
                self.txt_folder.setText(self.image_file_path[0])
            else:
                self.txt_folder.setText(os.path.dirname(self.image_file_path[0]))
            # 批量改名範圍同步
            self.rename_from.setText("1")
            self.rename_to.setText(str(count))

    def _on_clear_input(self):
        """清除全部照片。"""
        self.card_container.set_images([], [])
        self.image_file_path = []
        self.image_file_name_noext = []
        self.custom_names = []
        self._reset_input_state()

    def _reset_input_state(self):
        """清空照片後重置左側 UI 狀態。"""
        self.txt_folder.setText("")
        self.lbl_count.setVisible(False)
        self.btn_clear_input.setVisible(False)
        # 收起預覽面板並縮回視窗
        if self._preview_expanded:
            self.preview_frame.setVisible(False)
            self.resize(self.width() - 332, self.height())
            self._preview_expanded = False

    def _on_batch_rename(self):
        """批量改名：把第 from~to 張的名稱改成 新名稱+流水號。"""
        new_name = self.rename_new_name.text().strip()
        if not new_name:
            QMessageBox.warning(self, "未填新名稱", "請輸入新名稱後再套用。")
            return

        total = len(self.custom_names)
        if total == 0:
            return

        try:
            idx_from = int(self.rename_from.text()) - 1  # 轉 0-based
            idx_to = int(self.rename_to.text()) - 1
        except ValueError:
            QMessageBox.warning(self, "範圍格式錯誤", "請輸入正整數。")
            return

        # 邊界修正
        idx_from = max(0, min(idx_from, total - 1))
        idx_to = max(0, min(idx_to, total - 1))
        if idx_from > idx_to:
            idx_from, idx_to = idx_to, idx_from

        count = idx_to - idx_from + 1
        width = max(2, len(str(count)))  # 至少 2 位補齊

        cards = self.card_container._cards
        for seq, card_idx in enumerate(range(idx_from, idx_to + 1), start=1):
            if count == 1:
                name = new_name
            else:
                name = f"{new_name}{str(seq).zfill(width)}"
            # 更新 custom_names
            self.custom_names[card_idx] = name
            # 更新卡片顯示
            card = cards[card_idx]
            card.name = name
            card.name_label.setText(name)
            card.name_edit.setText(name)

    def _on_browse_output(self):
        start = self.path_output or self.path_templates
        folder = QFileDialog.getExistingDirectory(self, "選擇輸出資料夾", start)
        if folder:
            self.path_output = folder
            self.txt_output.setText(folder)

    def _on_run(self):
        title = self.input_title.text().strip()
        if not title:
            QMessageBox.warning(self, "未填標題", "請輸入頁首標題後再執行。")
            return
        if not self.image_file_path:
            QMessageBox.warning(self, "未選擇圖片", "請先選擇照片。")
            return
        if not self.path_output:
            QMessageBox.warning(self, "未選擇輸出資料夾", "請先選擇輸出資料夾。")
            return

        self.btn_run.setEnabled(False)
        self.log_box.clear()
        self._append_log(
            f"> 共 {len(self.image_file_path)} 張，模式：{self._selected_mode}"
        )
        self._append_log("> 開始處理...")

        self._worker = WorkerThread(
            mode=self._selected_mode,
            title=title,
            path_templates=self.path_templates,
            path_output=self.path_output,
            image_file_path=self.image_file_path,
            image_file_name_noext=self.custom_names,
        )
        self._worker.log.connect(self._append_log)
        self._worker.success.connect(self._on_success)
        self._worker.error.connect(self._on_error)
        self._worker.start()

    def _append_log(self, msg):
        self.log_box.append(msg)

    def _on_success(self, output_path):
        self._append_log(f"> 完成！輸出：{output_path}")
        self.btn_run.setEnabled(True)
        QMessageBox.information(self, "完成", f"已產生：\n{output_path}")

    def _on_error(self, tb):
        self._append_log("> [ERROR]\n" + tb)
        self.btn_run.setEnabled(True)
        QMessageBox.critical(self, "發生錯誤", tb)


# ── 入口 ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec())
